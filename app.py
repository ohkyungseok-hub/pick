# app.py (Streamlit Cloud 최종 통합본 - 오타/버그 수정 완료)
# ------------------------------------------------------------
# - 기본 로직/컬럼 유지
# - (옵션) 페이지 나누기 포함/제거 버전 둘 다 제공
# - XLSX/DOCX/PDF 모두 "페이지 번호 항상 표시"
# - DOCX: 표(행)가 페이지 경계에서 쪼개지지 않게(cantSplit)
# - DOCX: exact 고정 높이로 인한 "행 잘림" 방지 -> atLeast(최소 26pt)로 자동 확장
# - ✅ PDF: ReportLab로 직접 생성(한글 폰트 프로젝트 포함 방식)
# - ✅ Streamlit Cloud에서 PDF 미리보기: st.pdf() 사용(iframe/base64 방식 X)
#
# 실행:
#   streamlit run app.py
#
# requirements.txt 예시:
#   streamlit>=1.28
#   pandas
#   openpyxl
#   python-docx
#   reportlab
#
# 폰트 배치(중요: Streamlit Cloud):
#   repo_root/
#     app.py
#     requirements.txt
#     fonts/
#       NanumGothic.ttf
#       NanumGothicBold.ttf   (선택)
# ------------------------------------------------------------

import re
import tempfile
from pathlib import Path

import pandas as pd
import streamlit as st
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font
from openpyxl.worksheet.pagebreak import Break
from openpyxl.utils import get_column_letter

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

# ---- PDF (ReportLab)
REPORTLAB_OK = True
REPORTLAB_ERR = ""
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ModuleNotFoundError as e:
    REPORTLAB_OK = False
    REPORTLAB_ERR = str(e)


# -----------------------
# 컬럼 자동 매칭 유틸
# -----------------------
def norm_header(s: str) -> str:
    if s is None:
        return ""
    s = str(s).strip().lower()
    s = re.sub(r"\s+", "", s)
    s = re.sub(r"[()\-_/.,·:]", "", s)
    return s

AUTO_HEADER_CANDIDATES = {
    "상품연동코드": [
        "상품연동코드", "상품관리코드", "판매자관리코드", "판매자 관리코드",
        "상품코드", "품목코드", "sku", "sku번호"
    ],
    "주문상품": [
        "주문상품", "상품명", "주문상품명", "노출상품명", "노출상품명(옵션명)"
    ],
    "옵션": [
        "옵션", "옵션명", "옵션정보", "옵션명:옵션값", "옵션값"
    ],
    "주문수량": [
        "주문수량", "수량", "구매수", "구매수량", "qty"
    ],
    "수령자": [
        "수령자", "수령자명", "수취인", "수취인명", "받는분", "받는분성명"
    ],
    "주소": [
        "주소", "배송지주소", "수취인주소", "받는분주소", "받는분주소(전체,분할)"
    ],
    "주문요청사항": [
        "주문요청사항", "주문메모", "배송메모", "배송요청사항", "배송메시지", "배송메세지"
    ],
}

KNOWN_IMPORT_FORMATS = [
    {
        "name": "결제완료 리스트",
        "required_keys": ["상품연동코드", "주문상품", "옵션", "주문수량", "수령자", "주소"],
        "columns": {
            "상품연동코드": "R",
            "주문상품": "T",
            "옵션": "X",
            "주문수량": "AD",
            "수령자": "AM",
            "주소": "AP",
            "주문요청사항": "F",
        },
        "expected_headers": {
            "상품연동코드": "상품관리코드",
            "주문상품": "상품명",
            "옵션": "옵션명:옵션값",
            "주문수량": "수량",
            "수령자": "수령자명",
            "주소": "주소",
        },
    },
]

def find_header_col(df: pd.DataFrame, candidates: list[str]):
    norm_map = {norm_header(c): c for c in df.columns}

    # 완전 일치
    for cand in candidates:
        nc = norm_header(cand)
        if nc in norm_map:
            return norm_map[nc]

    # 부분 포함
    for nc_df, original in norm_map.items():
        for cand in candidates:
            nc = norm_header(cand)
            if nc and (nc in nc_df or nc_df in nc):
                return original

    return None

def get_df_col_by_excel_letter(df: pd.DataFrame, col_letter: str):
    idx = excel_col_to_zero_index(col_letter)
    if idx >= df.shape[1]:
        return None
    return df.columns[idx]

def detect_known_import_format(df: pd.DataFrame) -> tuple[dict, str | None]:
    for fmt in KNOWN_IMPORT_FORMATS:
        expected_headers = fmt["expected_headers"]
        matched = True
        for key in fmt["required_keys"]:
            source_col = get_df_col_by_excel_letter(df, fmt["columns"][key])
            if source_col is None:
                matched = False
                break
            if norm_header(source_col) != norm_header(expected_headers[key]):
                matched = False
                break

        if matched:
            detected = {}
            for key, col_letter in fmt["columns"].items():
                source_col = get_df_col_by_excel_letter(df, col_letter)
                if source_col is not None:
                    detected[key] = source_col
            return detected, fmt["name"]

    return {}, None

def auto_detect_colmap(df: pd.DataFrame) -> dict:
    detected = {}
    for key, candidates in AUTO_HEADER_CANDIDATES.items():
        detected[key] = find_header_col(df, candidates)
    return detected


# -----------------------
# 기본 로직(원본과 동일)
# -----------------------
def excel_col_to_zero_index(col_letter: str) -> int:
    """Excel column letter (e.g., 'A', 'J') -> pandas zero-based index"""
    col_letter = col_letter.strip().upper()
    n = 0
    for ch in col_letter:
        if not ("A" <= ch <= "Z"):
            raise ValueError(f"Invalid column letter: {col_letter}")
        n = n * 26 + (ord(ch) - ord("A") + 1)
    return n - 1


def build_picking_dataframe(src_path: str, colmap: dict | None = None) -> tuple[pd.DataFrame, dict]:
    """원본 엑셀 -> 피킹용 DF(주소 정렬 + 주소별 합계행 포함)"""
    df = pd.read_excel(src_path)

    need_keys = ["상품연동코드", "주문상품", "옵션", "주문수량", "수령자", "주소", "주문요청사항"]
    optional_keys = {"주문요청사항"}

    # 1) 헤더 자동 탐지
    detected_map = auto_detect_colmap(df)
    known_format_map, known_format_name = detect_known_import_format(df)

    # 2) 수동 입력(colmap)이 있으면 수동값 우선, 없으면 자동값/알려진 양식값 사용
    final_cols = {}
    for k in need_keys:
        user_val = (colmap or {}).get(k, "")
        user_val = str(user_val).strip() if user_val is not None else ""

        if user_val:
            idx = excel_col_to_zero_index(user_val)
            if idx >= df.shape[1]:
                raise ValueError(f"{k}: 입력한 열 {user_val} 이 실제 컬럼 범위를 벗어났습니다.")
            final_cols[k] = df.columns[idx]
        else:
            final_cols[k] = detected_map.get(k) or known_format_map.get(k)

    missing = [k for k in need_keys if k not in optional_keys and not final_cols.get(k)]
    if missing:
        raise ValueError(
            f"자동 매칭 실패 컬럼: {missing}\n"
            f"현재 파일 헤더: {list(df.columns)}"
        )

    df_sel = pd.DataFrame()
    for k in need_keys:
        source_col = final_cols.get(k)
        if source_col:
            df_sel[k] = df[source_col]
        else:
            df_sel[k] = ""
            final_cols[k] = "(빈 값)"

    if known_format_name:
        final_cols["_인식양식"] = known_format_name

    # 정렬: 주소(오름), 상품연동코드(내림)
    df_sorted = df_sel.sort_values(
        by=["주소", "상품연동코드"],
        ascending=[True, False],
        kind="mergesort",
    )

    # 주소별 합계행 추가
    out_chunks = []
    for addr, g in df_sorted.groupby("주소", sort=False):
        out_chunks.append(g)

        subtotal = {c: "" for c in df_sorted.columns}
        subtotal["주문상품"] = "합계"
        subtotal_qty = pd.to_numeric(g["주문수량"], errors="coerce").fillna(0).sum()
        subtotal["주문수량"] = subtotal_qty
        subtotal["주소"] = addr
        out_chunks.append(pd.DataFrame([subtotal]))

    df_final = pd.concat(out_chunks, ignore_index=True)
    return df_final, final_cols


def build_picking_xlsx(df_final: pd.DataFrame, out_path: str, add_page_breaks: bool = True) -> None:
    """
    DF -> 피킹용 엑셀 저장 + 인쇄/서식 설정(openpyxl)
    add_page_breaks=True : 주소 바뀔 때마다 페이지 나누기
    add_page_breaks=False: 페이지 나누기 제거
    + 페이지 번호(하단 중앙) 항상 표시
    """
    df_final.to_excel(out_path, index=False)

    wb = load_workbook(out_path)
    ws = wb.active

    # 헤더: 굵게 + 줄바꿈
    header_font = Font(bold=True)
    header_align = Alignment(wrap_text=True, vertical="center")
    for c in range(1, ws.max_column + 1):
        cell = ws.cell(1, c)
        cell.font = header_font
        cell.alignment = header_align

    headers = {ws.cell(row=1, column=c).value: c for c in range(1, ws.max_column + 1)}
    addr_col = headers["주소"]

    # 긴 텍스트 줄바꿈 + 위쪽 정렬
    wrap_top = Alignment(wrap_text=True, vertical="top")
    for r in range(2, ws.max_row + 1):
        for name in ["주문상품", "옵션", "주소", "주문요청사항"]:
            ws.cell(r, headers[name]).alignment = wrap_top

    # 열 너비(기존 유지)
    widths = {
        "상품연동코드": 18,
        "주문상품": 60,
        "옵션": 50,
        "주문수량": 10,
        "수령자": 18,
        "주소": 50,
        "주문요청사항": 40,
    }
    for name, w in widths.items():
        ws.column_dimensions[get_column_letter(headers[name])].width = w

    # 인쇄 설정(기존 유지: 가로)
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_title_rows = "1:1"

    # 페이지 번호 항상 표시(하단 중앙)
    ws.oddFooter.center.text = "페이지 &P / &N"
    ws.evenFooter.center.text = "페이지 &P / &N"

    # 주소 바뀔 때마다 페이지 나누기(옵션)
    ws.row_breaks.brk = []
    if add_page_breaks and ws.max_row >= 2:
        prev_addr = ws.cell(2, addr_col).value
        for r in range(3, ws.max_row + 1):
            curr_addr = ws.cell(r, addr_col).value
            if curr_addr != prev_addr:
                ws.row_breaks.append(Break(id=r - 1))  # 이전 행 뒤에서 끊기
                prev_addr = curr_addr

    ws.print_area = f"A1:{get_column_letter(ws.max_column)}{ws.max_row}"
    wb.save(out_path)


# -----------------------
# 워드 출력 (세로모드 최종)
# -----------------------
def add_page_number_footer(section) -> None:
    """워드 하단 중앙에 페이지 번호(PAGE 필드) 삽입"""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = 1  # CENTER

    # 기존 내용이 있으면 덧붙지 않도록 초기화
    for r in list(p.runs):
        r.text = ""

    run = p.add_run()

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _docx_set_row_layout(row, min_pt: int = 26) -> None:
    """
    - 행이 페이지 경계에서 쪼개지지 않게(cantSplit)
    - 행 높이: exact가 아니라 atLeast(최소 높이)로 설정 -> 내용 길면 자동 확장(잘림 방지)
    """
    trPr = row._tr.get_or_add_trPr()

    # 행 분할 금지
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)

    # 최소 높이만 지정(내용에 따라 자동 확장)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(min_pt * 20)))  # twips
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)

    # 문단 여백 제거
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0


def _docx_shade_row(row, fill: str = "EFEFEF") -> None:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)


def build_picking_docx(df_final: pd.DataFrame, out_docx: str, add_page_breaks: bool = True) -> None:
    """
    DF -> 피킹용 워드(.docx)
    - A4 세로
    - 행높이: 최소 26pt(내용에 따라 자동 확장) -> 행 잘림 방지
    - 주소별로 표 생성
    - add_page_breaks=True : 주소별 1페이지(페이지 나누기)
    - add_page_breaks=False: 페이지 나누기 제거(연속 출력)
    - 코드 변경 시 음영 토글
    - 페이지 번호 항상 표시(하단 중앙)
    """
    required_cols = ["상품연동코드", "주문상품", "옵션", "주문수량", "수령자", "주소", "주문요청사항"]
    for c in required_cols:
        if c not in df_final.columns:
            raise ValueError(f"df_final에 '{c}' 컬럼이 없습니다. 현재 컬럼: {list(df_final.columns)}")

    doc = Document()

    # 페이지(세로) + 여백
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69) # A4
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    # 페이지 번호
    add_page_number_footer(section)

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(9)

    # 주소별 그룹 만들기(기존 유지)
    groups = []
    current_addr = None
    current_rows = []
    for _, row in df_final.iterrows():
        addr = "" if pd.isna(row["주소"]) else str(row["주소"]).strip()
        if current_addr is None:
            current_addr = addr
            current_rows = [row]
        elif addr == current_addr:
            current_rows.append(row)
        else:
            groups.append((current_addr, current_rows))
            current_addr = addr
            current_rows = [row]
    if current_rows:
        groups.append((current_addr, current_rows))

    cols = required_cols[:]

    # 세로모드 열 너비(인치)
    col_widths = {
        "상품연동코드": Inches(0.8),
        "주문상품": Inches(2.4),
        "옵션": Inches(1.4),
        "주문수량": Inches(0.6),
        "수령자": Inches(1.0),
        "주소": Inches(1.2),
        "주문요청사항": Inches(1.1),
    }

    for gi, (addr, rows_in_addr) in enumerate(groups):
        # 상단 주소(10pt) + 표와 붙이기
        p = doc.add_paragraph(f"주소: {addr}")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True

        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False

        # 헤더행
        hdr = table.rows[0]
        _docx_set_row_layout(hdr, 26)
        for ci, name in enumerate(cols):
            cell = hdr.cells[ci]
            cell.text = name
            cell.width = col_widths[name]
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(8)

        # 코드 변경 시 음영 토글
        last_code = None
        shade_on = False

        for r in rows_in_addr:
            is_sum = (str(r.get("주문상품", "")) == "합계") or ("합계" in str(r.get("주문상품", "")))

            code_val = "" if pd.isna(r["상품연동코드"]) else str(r["상품연동코드"])
            if not is_sum and code_val != last_code:
                shade_on = not shade_on
                last_code = code_val

            row = table.add_row()
            _docx_set_row_layout(row, 26)

            if shade_on and not is_sum:
                _docx_shade_row(row, "EFEFEF")

            for ci, name in enumerate(cols):
                cell = row.cells[ci]
                cell.width = col_widths[name]

                # 합계행: 주소칸 비움
                if is_sum and name == "주소":
                    cell.text = ""
                    continue

                cell.text = ""
                val = r.get(name, "")
                text = "" if pd.isna(val) else str(val)
                run = cell.paragraphs[0].add_run(text)

                # 폰트 규칙
                if name == "주소":
                    run.font.size = Pt(5)
                elif name in ("주문상품", "옵션"):
                    run.font.size = Pt(8)
                elif name == "주문수량":
                    run.font.size = Pt(12)
                    try:
                        if int(float(text)) >= 2:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                    except Exception:
                        pass
                elif name == "상품연동코드":
                    run.font.size = Pt(14)
                    run.bold = True
                else:
                    run.font.size = Pt(8)

                if is_sum:
                    run.font.size = Pt(16)
                    run.bold = True

        # 페이지 나누기(옵션)
        if add_page_breaks and gi != len(groups) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif (not add_page_breaks) and gi != len(groups) - 1:
            doc.add_paragraph("")

    doc.save(out_docx)


# -----------------------
# PDF 출력 (Streamlit Cloud: 프로젝트 폰트 사용)
# -----------------------
def _register_korean_font_from_project() -> tuple[str, str, bool]:
    """
    프로젝트 내 ./fonts 폴더의 폰트를 PDF에 등록하여 한글(■) 깨짐 방지.
    - repo_root/fonts/NanumGothic.ttf 필수
    - repo_root/fonts/NanumGothicBold.ttf 선택
    반환: (regular_font_name, bold_font_name, font_ok)
    """
    base_dir = Path(__file__).parent
    regular_path = base_dir / "fonts" / "NanumGothic.ttf"
    bold_path = base_dir / "fonts" / "NanumGothicBold.ttf"

    if not regular_path.exists():
        return "Helvetica", "Helvetica-Bold", False

    pdfmetrics.registerFont(TTFont("KOREAN", str(regular_path)))
    if bold_path.exists():
        pdfmetrics.registerFont(TTFont("KOREAN_B", str(bold_path)))
    else:
        pdfmetrics.registerFont(TTFont("KOREAN_B", str(regular_path)))

    return "KOREAN", "KOREAN_B", True


def build_picking_pdf(df_final: pd.DataFrame, out_pdf: str, add_page_breaks: bool = False) -> tuple[str, bool]:
    """
    DF -> 피킹용 PDF
    - A4 세로
    - 페이지 번호(하단 중앙) 항상 표시
    - 주소별로 블록 구성
    - add_page_breaks=True  : 주소별 페이지 강제 분리(PageBreak)
      add_page_breaks=False : 강제 분리 없이 연속 출력(자연 페이지네이션) ✅(페이지 나누기 제거 목적)
    - 표 행 분할 방지: NOSPLIT 적용
    반환: (font_used, font_ok)
    """
    required_cols = ["상품연동코드", "주문상품", "옵션", "주문수량", "수령자", "주소", "주문요청사항"]
    for c in required_cols:
        if c not in df_final.columns:
            raise ValueError(f"df_final에 '{c}' 컬럼이 없습니다. 현재 컬럼: {list(df_final.columns)}")

    # 주소별 그룹 (DOCX 로직과 동일)
    groups = []
    current_addr = None
    current_rows = []
    for _, row in df_final.iterrows():
        addr = "" if pd.isna(row["주소"]) else str(row["주소"]).strip()
        if current_addr is None:
            current_addr = addr
            current_rows = [row]
        elif addr == current_addr:
            current_rows.append(row)
        else:
            groups.append((current_addr, current_rows))
            current_addr = addr
            current_rows = [row]
    if current_rows:
        groups.append((current_addr, current_rows))

    font_regular, font_bold, font_ok = _register_korean_font_from_project()

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "base",
        parent=styles["Normal"],
        fontName=font_regular,
        fontSize=8,
        leading=10,
    )
    title_style = ParagraphStyle(
        "title",
        parent=base,
        fontName=font_bold,
        fontSize=10,
        leading=12,
        spaceAfter=4,
    )

    def draw_page_number(canvas, doc):
        page_num = canvas.getPageNumber()
        canvas.saveState()
        canvas.setFont(font_regular, 9)
        canvas.drawCentredString(A4[0] / 2.0, 10 * mm, f"{page_num}")
        canvas.restoreState()

    pdf_doc = SimpleDocTemplate(
        out_pdf,
        pagesize=A4,
        leftMargin=9 * mm,
        rightMargin=9 * mm,
        topMargin=9 * mm,
        bottomMargin=12 * mm,
    )

    flow = []
    cols = required_cols[:]

    # A4 세로 폭에 맞춘 컬럼 폭(mm)
    col_widths = [18 * mm, 58 * mm, 34 * mm, 14 * mm, 22 * mm, 28 * mm, 26 * mm]

    for gi, (addr, rows_in_addr) in enumerate(groups):
        flow.append(Paragraph(f"주소: {addr}", title_style))

        # 표 데이터(Paragraph로 넣어 줄바꿈/자동높이)
        data = []
        data.append([Paragraph(str(c), base) for c in cols])

        last_code = None
        shade_on = False
        row_shaded = [False]  # header

        for r in rows_in_addr:
            is_sum = (str(r.get("주문상품", "")) == "합계") or ("합계" in str(r.get("주문상품", "")))
            code_val = "" if pd.isna(r["상품연동코드"]) else str(r["상품연동코드"])

            if not is_sum and code_val != last_code:
                shade_on = not shade_on
                last_code = code_val

            row_vals = []
            for name in cols:
                if is_sum and name == "주소":
                    row_vals.append("")
                else:
                    v = r.get(name, "")
                    row_vals.append("" if pd.isna(v) else str(v))

            data.append([Paragraph(v.replace("\n", "<br/>"), base) for v in row_vals])
            row_shaded.append(bool(shade_on and (not is_sum)))

        t = Table(data, colWidths=col_widths, repeatRows=1)

        ts = TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.25, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),

            ("FONTNAME", (0, 0), (-1, -1), font_regular),
            ("FONTSIZE", (0, 0), (-1, -1), 8),

            ("FONTNAME", (0, 0), (-1, 0), font_bold),
            ("FONTSIZE", (0, 0), (-1, 0), 8),

            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("ALIGN", (3, 1), (3, -1), "CENTER"),

            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ])

        # 행 분할 방지: 각 행마다 NOSPLIT
        for ri in range(0, len(data)):
            ts.add("NOSPLIT", (0, ri), (-1, ri))

        # 코드 변경 음영 토글
        for ri in range(1, len(data)):
            if row_shaded[ri]:
                ts.add("BACKGROUND", (0, ri), (-1, ri), colors.whitesmoke)

        t.setStyle(ts)
        flow.append(t)
        flow.append(Spacer(1, 6))

        if add_page_breaks and gi != len(groups) - 1:
            flow.append(PageBreak())

    pdf_doc.build(flow, onFirstPage=draw_page_number, onLaterPages=draw_page_number)
    return font_regular, font_ok


# -----------------------
# Streamlit UI
# -----------------------
st.set_page_config(page_title="피킹 시트 생성기", layout="wide")
st.title("피킹 시트 생성기 (Excel → Picking XLSX / DOCX / PDF)")

st.write(
    "- 원본 엑셀 업로드 → **주소별 정렬 + 주소별 합계행**\n"
    "- **XLSX / DOCX / PDF** 생성\n"
    "- ✅ 페이지 번호 항상 표시(XLSX/DOCX/PDF)\n"
    "- ✅ DOCX 표 행은 페이지에서 쪼개지지 않음(안 들어가면 다음 페이지로 이동)\n"
    "- ✅ DOCX 행 잘림 방지(최소 26pt + 자동 확장)\n"
    "- (옵션) **페이지 나누기 제거 버전**도 함께 생성\n"
    "- ✅ PDF는 Streamlit Cloud에서 **st.pdf**로 바로 미리보기"
)

uploaded = st.file_uploader("원본 엑셀 업로드 (.xlsx)", type=["xlsx"])

with st.expander("원본 컬럼 매핑(비워두면 헤더명으로 자동 매칭)"):
    st.caption("비워두면 헤더명으로 자동 매칭합니다. 자동 매칭이 안 될 때만 열(letter)을 입력하세요. 예: J")
    default_map = {
        "상품연동코드": "",
        "주문상품": "",
        "옵션": "",
        "주문수량": "",
        "수령자": "",
        "주소": "",
        "주문요청사항": "",
    }
    colmap = {}
    cols_ui = st.columns(len(default_map))
    keys = list(default_map.keys())
    for i, k in enumerate(keys):
        with cols_ui[i]:
            colmap[k] = st.text_input(k, value=default_map[k], max_chars=3)

make_xlsx = st.checkbox("결과 XLSX 생성", value=True)
make_docx = st.checkbox("결과 DOCX 생성", value=True)

if not REPORTLAB_OK:
    st.warning(
        "PDF 기능을 사용하려면 reportlab 설치가 필요합니다. requirements.txt에 reportlab을 추가하세요.\n"
        f"(현재 오류: {REPORTLAB_ERR})"
    )

make_pdf = st.checkbox("결과 PDF 생성(페이지 나누기 제거 목적)", value=True, disabled=not REPORTLAB_OK)
preview_pdf = st.checkbox("PDF를 화면에서 바로 미리보기", value=True, disabled=not REPORTLAB_OK)

also_make_no_pagebreak = st.checkbox("페이지 나누기 제거 버전도 함께 생성", value=True)

base_name = st.text_input("파일명 접두어(다운로드 파일명)", value="picking_result")
run_btn = st.button("생성하기", type="primary", disabled=(uploaded is None))

if run_btn:
    if uploaded is None:
        st.error("원본 엑셀 파일을 업로드하세요.")
        st.stop()

    if not (make_xlsx or make_docx or make_pdf):
        st.warning("XLSX/DOCX/PDF 중 최소 1개는 선택해야 합니다.")
        st.stop()

    try:
        with tempfile.TemporaryDirectory() as td:
            td_path = Path(td)
            src_path = td_path / "source.xlsx"
            src_path.write_bytes(uploaded.getvalue())

            # 1) DF 생성
            df_final, detected_cols = build_picking_dataframe(str(src_path), colmap)

            st.success("데이터 변환 완료! (주소별 정렬 + 합계행 생성)")
            st.caption("자동/최종 매칭 컬럼")
            st.json(detected_cols)
            st.dataframe(df_final, use_container_width=True, height=360)

            # (진단) 폰트 파일 존재 여부 표시(Cloud에서 유용)
            font_path = Path(__file__).parent / "fonts" / "NanumGothic.ttf"
            if make_pdf and REPORTLAB_OK:
                st.caption(f"폰트 경로 확인: {font_path} / exists={font_path.exists()}")

            # 2) XLSX 생성(기본: 페이지 나누기 포함)
            if make_xlsx:
                out_xlsx_path = td_path / f"{base_name}.xlsx"
                build_picking_xlsx(df_final, str(out_xlsx_path), add_page_breaks=True)
                st.download_button(
                    label="📥 결과 XLSX 다운로드 (페이지 나누기 포함)",
                    data=out_xlsx_path.read_bytes(),
                    file_name=f"{base_name}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

                if also_make_no_pagebreak:
                    out_xlsx_np_path = td_path / f"{base_name}_nopagebreak.xlsx"
                    build_picking_xlsx(df_final, str(out_xlsx_np_path), add_page_breaks=False)
                    st.download_button(
                        label="📥 결과 XLSX 다운로드 (페이지 나누기 제거)",
                        data=out_xlsx_np_path.read_bytes(),
                        file_name=f"{base_name}_nopagebreak.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )

            # 3) DOCX 생성(기본: 주소별 페이지 나누기 포함)
            if make_docx:
                out_docx_path = td_path / f"{base_name}.docx"
                build_picking_docx(df_final, str(out_docx_path), add_page_breaks=True)
                st.download_button(
                    label="📥 결과 DOCX 다운로드 (주소별 페이지 나누기 포함)",
                    data=out_docx_path.read_bytes(),
                    file_name=f"{base_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

                if also_make_no_pagebreak:
                    out_docx_np_path = td_path / f"{base_name}_nopagebreak.docx"
                    build_picking_docx(df_final, str(out_docx_np_path), add_page_breaks=False)
                    st.download_button(
                        label="📥 결과 DOCX 다운로드 (페이지 나누기 제거)",
                        data=out_docx_np_path.read_bytes(),
                        file_name=f"{base_name}_nopagebreak.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

            # 4) PDF 생성(페이지 나누기 제거 목적)
            if make_pdf and REPORTLAB_OK:
                out_pdf_np_path = td_path / f"{base_name}_nopagebreak.pdf"
                font_used, font_ok = build_picking_pdf(df_final, str(out_pdf_np_path), add_page_breaks=False)
                pdf_bytes = out_pdf_np_path.read_bytes()

                st.download_button(
                    label="📥 결과 PDF 다운로드 (페이지 나누기 제거/자연 페이지네이션)",
                    data=pdf_bytes,
                    file_name=f"{base_name}_nopagebreak.pdf",
                    mime="application/pdf",
                )

                if not font_ok:
                    st.warning(
                        "PDF 한글 폰트 등록에 실패해 기본 폰트로 생성되었습니다(한글이 ■로 깨질 수 있어요). "
                        "repo_root/fonts/NanumGothic.ttf가 있는지 확인하세요."
                    )
                else:
                    st.caption(f"PDF 폰트 적용: {font_used}")

                if preview_pdf:
                    st.subheader("PDF 미리보기")
                    # Streamlit Cloud에서 iframe/base64 대신 st.pdf 사용
                    st.pdf(pdf_bytes)

    except Exception as e:
        st.error("생성 중 오류가 발생했습니다.")
        st.exception(e)
